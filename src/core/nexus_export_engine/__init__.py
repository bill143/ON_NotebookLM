"""
Nexus Export Engine — Multi-Format Document Export
Codename: ESPERANTO — Feature 6B: Export & Sharing

Provides:
- PDF export with styled layout and table of contents
- DOCX export with heading hierarchy and tables
- EPUB export for e-reader consumption
- Markdown export (native)
- Templated export with header/footer branding
- Batch export for multiple artifacts
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import UTC, datetime
from typing import Any

from loguru import logger

from src.exceptions import ValidationError
from src.infra.nexus_obs_tracing import traced

# ── Types ────────────────────────────────────────────────────


class ExportFormat:
    PDF = "pdf"
    DOCX = "docx"
    EPUB = "epub"
    MARKDOWN = "markdown"
    HTML = "html"
    TXT = "txt"


@dataclass
class ExportContent:
    """Content to be exported."""

    title: str
    content: str  # Markdown-formatted content
    author: str = "Nexus Notebook 11 LM"
    created_at: str | None = None
    notebook_name: str = ""
    content_type: str = ""  # "summary", "quiz", "report", etc.
    sections: list[dict[str, str]] = field(default_factory=list)  # {title, content}
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ExportOptions:
    """Export configuration."""

    format: str = ExportFormat.PDF
    include_toc: bool = True
    include_header: bool = True
    include_footer: bool = True
    include_metadata: bool = True
    page_size: str = "A4"  # "A4", "Letter"
    font_family: str = "Helvetica"
    font_size: int = 11
    line_spacing: float = 1.4
    margin_mm: int = 25
    branding_color: str = "#6366f1"  # Primary brand color
    watermark: str | None = None


@dataclass
class ExportResult:
    """Result of an export operation."""

    data: bytes
    filename: str
    mime_type: str
    format: str
    file_size_bytes: int


# ── Markdown Parser ──────────────────────────────────────────


class MarkdownParser:
    """Parse markdown into structured sections for export."""

    @staticmethod
    def parse_sections(content: str) -> list[dict[str, str]]:
        """Split markdown into sections by headers."""
        sections: list[dict[str, str]] = []
        current_title = ""
        current_body: list[str] = []

        for line in content.split("\n"):
            if line.startswith("# "):
                if current_body:
                    sections.append(
                        {
                            "title": current_title,
                            "content": "\n".join(current_body).strip(),
                            "level": "1",
                        }
                    )
                current_title = line[2:].strip()
                current_body = []
            elif line.startswith("## "):
                if current_body:
                    sections.append(
                        {
                            "title": current_title,
                            "content": "\n".join(current_body).strip(),
                            "level": "2" if current_title else "1",
                        }
                    )
                current_title = line[3:].strip()
                current_body = []
            elif line.startswith("### "):
                if current_body:
                    sections.append(
                        {
                            "title": current_title,
                            "content": "\n".join(current_body).strip(),
                            "level": "3",
                        }
                    )
                current_title = line[4:].strip()
                current_body = []
            else:
                current_body.append(line)

        if current_body:
            sections.append(
                {
                    "title": current_title,
                    "content": "\n".join(current_body).strip(),
                    "level": "1",
                }
            )

        return sections

    @staticmethod
    def strip_markdown(text: str) -> str:
        """Remove markdown formatting for plain text."""
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
        text = re.sub(r"#{1,6}\s+", "", text)
        text = re.sub(r"^[-*+]\s+", "• ", text, flags=re.MULTILINE)
        return text

    @staticmethod
    def to_html(content: str) -> str:
        """Convert markdown to HTML for EPUB/PDF."""
        try:
            import markdown

            return markdown.markdown(
                content,
                extensions=["tables", "fenced_code", "toc", "nl2br"],
            )
        except ImportError:
            # Fallback: basic conversion
            html = content
            html = re.sub(r"^### (.+)$", r"<h3>\1</h3>", html, flags=re.MULTILINE)
            html = re.sub(r"^## (.+)$", r"<h2>\1</h2>", html, flags=re.MULTILINE)
            html = re.sub(r"^# (.+)$", r"<h1>\1</h1>", html, flags=re.MULTILINE)
            html = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", html)
            html = re.sub(r"\*(.+?)\*", r"<em>\1</em>", html)
            html = re.sub(r"`(.+?)`", r"<code>\1</code>", html)
            html = re.sub(r"\n\n", "</p><p>", html)
            return f"<p>{html}</p>"


# ── PDF Exporter ─────────────────────────────────────────────


class PDFExporter:
    """Export content to PDF using reportlab."""

    @traced("export.pdf")
    async def export(
        self,
        content: ExportContent,
        options: ExportOptions,
    ) -> ExportResult:
        """Generate a styled PDF document."""
        try:
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
            from reportlab.lib.pagesizes import A4, letter
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.platypus import (
                HRFlowable,
                Paragraph,
                SimpleDocTemplate,
                Spacer,
            )
        except ImportError as e:
            raise ValidationError(
                "reportlab is required for PDF export. Install with: pip install reportlab"
            ) from e

        buffer = io.BytesIO()
        page_size = A4 if options.page_size == "A4" else letter
        margin = options.margin_mm * mm

        doc = SimpleDocTemplate(
            buffer,
            pagesize=page_size,
            leftMargin=margin,
            rightMargin=margin,
            topMargin=margin,
            bottomMargin=margin,
            title=content.title,
            author=content.author,
        )

        # Styles
        styles = getSampleStyleSheet()
        brand_color = colors.HexColor(options.branding_color)

        title_style = ParagraphStyle(
            "NexusTitle",
            parent=styles["Title"],
            fontSize=24,
            spaceAfter=12,
            textColor=brand_color,
            fontName=options.font_family,
        )

        heading_style = ParagraphStyle(
            "NexusH2",
            parent=styles["Heading2"],
            fontSize=16,
            spaceBefore=18,
            spaceAfter=8,
            textColor=brand_color,
            fontName=options.font_family,
        )

        h3_style = ParagraphStyle(
            "NexusH3",
            parent=styles["Heading3"],
            fontSize=13,
            spaceBefore=12,
            spaceAfter=6,
            textColor=colors.HexColor("#374151"),
            fontName=options.font_family,
        )

        body_style = ParagraphStyle(
            "NexusBody",
            parent=styles["Normal"],
            fontSize=options.font_size,
            leading=options.font_size * options.line_spacing,
            alignment=TA_JUSTIFY,
            fontName=options.font_family,
        )

        meta_style = ParagraphStyle(
            "NexusMeta",
            parent=styles["Normal"],
            fontSize=9,
            textColor=colors.HexColor("#9ca3af"),
            alignment=TA_CENTER,
            spaceAfter=20,
        )

        # Build story
        story: list = []

        # Title page
        story.append(Spacer(1, 60))
        story.append(Paragraph(content.title, title_style))

        if content.notebook_name:
            story.append(Paragraph(f"From: {content.notebook_name}", meta_style))

        story.append(
            Paragraph(
                f"Generated: {content.created_at or datetime.now(UTC).strftime('%B %d, %Y')}",
                meta_style,
            )
        )

        story.append(
            HRFlowable(
                width="80%",
                color=brand_color,
                thickness=2,
                spaceBefore=10,
                spaceAfter=20,
            )
        )

        # Parse and render content
        sections = content.sections or MarkdownParser.parse_sections(content.content)

        for section in sections:
            level = section.get("level", "1")
            title = section.get("title", "")
            body = section.get("content", "")

            if title:
                style = heading_style if level in ("1", "2") else h3_style
                story.append(Paragraph(title, style))

            # Process body text
            for paragraph in body.split("\n\n"):
                paragraph = paragraph.strip()
                if not paragraph:
                    continue

                # Handle bullet points
                if paragraph.startswith(("- ", "* ", "• ")):
                    for line in paragraph.split("\n"):
                        clean = re.sub(r"^[-*•]\s+", "", line)
                        if clean:
                            story.append(Paragraph(f"• {clean}", body_style))
                else:
                    # Clean markdown formatting
                    clean = MarkdownParser.strip_markdown(paragraph)
                    story.append(Paragraph(clean, body_style))

                story.append(Spacer(1, 4))

        # Footer
        if options.include_footer:
            story.append(Spacer(1, 30))
            story.append(
                HRFlowable(
                    width="60%",
                    color=colors.HexColor("#e5e7eb"),
                    thickness=0.5,
                    spaceBefore=10,
                    spaceAfter=10,
                )
            )
            story.append(
                Paragraph(
                    "Generated by Nexus Notebook 11 LM — Codename: ESPERANTO",
                    meta_style,
                )
            )

        doc.build(story)
        pdf_data = buffer.getvalue()

        filename = self._safe_filename(content.title) + ".pdf"

        return ExportResult(
            data=pdf_data,
            filename=filename,
            mime_type="application/pdf",
            format="pdf",
            file_size_bytes=len(pdf_data),
        )

    @staticmethod
    def _safe_filename(title: str) -> str:
        """Generate a safe filename from a title."""
        safe = re.sub(r"[^\w\s-]", "", title.lower())
        safe = re.sub(r"[-\s]+", "_", safe)
        return safe[:80]


# ── DOCX Exporter ────────────────────────────────────────────


class DOCXExporter:
    """Export content to DOCX using python-docx."""

    @traced("export.docx")
    async def export(
        self,
        content: ExportContent,
        options: ExportOptions,
    ) -> ExportResult:
        """Generate a styled DOCX document."""
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt, RGBColor
        except ImportError as e:
            raise ValidationError(
                "python-docx is required for DOCX export. Install with: pip install python-docx"
            ) from e

        doc = Document()

        # Document properties
        core = doc.core_properties
        core.title = content.title
        core.author = content.author

        # Style the default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(options.font_size)

        # Title
        title_para = doc.add_heading(content.title, level=0)
        for run in title_para.runs:
            run.font.color.rgb = RGBColor.from_string(options.branding_color.lstrip("#"))

        # Metadata
        if options.include_metadata:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = meta_para.add_run(
                f"Generated: {content.created_at or datetime.now(UTC).strftime('%B %d, %Y')}"
            )
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(156, 163, 175)

            if content.notebook_name:
                run2 = meta_para.add_run(f"\nNotebook: {content.notebook_name}")
                run2.font.size = Pt(9)
                run2.font.color.rgb = RGBColor(156, 163, 175)

        doc.add_paragraph("")  # Spacer

        # Render content
        sections = content.sections or MarkdownParser.parse_sections(content.content)

        for section in sections:
            level = int(section.get("level", "1"))
            title = section.get("title", "")
            body = section.get("content", "")

            if title:
                heading_level = min(level, 4)
                heading = doc.add_heading(title, level=heading_level)
                if level <= 2:
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor.from_string(
                            options.branding_color.lstrip("#")
                        )

            for paragraph in body.split("\n\n"):
                paragraph = paragraph.strip()
                if not paragraph:
                    continue

                if paragraph.startswith(("- ", "* ", "• ")):
                    for line in paragraph.split("\n"):
                        clean = re.sub(r"^[-*•]\s+", "", line)
                        if clean:
                            doc.add_paragraph(clean, style="List Bullet")
                else:
                    # Handle bold and italic within the text
                    para = doc.add_paragraph()
                    self._add_formatted_text(para, paragraph)

        # Footer
        if options.include_footer:
            doc.add_paragraph("")
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = footer_para.add_run("Generated by Nexus Notebook 11 LM — Codename: ESPERANTO")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(156, 163, 175)

        # Save
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_data = buffer.getvalue()

        filename = PDFExporter._safe_filename(content.title) + ".docx"

        return ExportResult(
            data=docx_data,
            filename=filename,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            format="docx",
            file_size_bytes=len(docx_data),
        )

    @staticmethod
    def _add_formatted_text(para, text: str):
        """Add text with basic bold/italic formatting."""
        # Simple bold/italic parsing
        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                run.italic = True
            else:
                clean = MarkdownParser.strip_markdown(part)
                para.add_run(clean)


# ── EPUB Exporter ────────────────────────────────────────────


class EPUBExporter:
    """Export content to EPUB for e-reader consumption."""

    @traced("export.epub")
    async def export(
        self,
        content: ExportContent,
        options: ExportOptions,
    ) -> ExportResult:
        """Generate an EPUB document."""
        try:
            from ebooklib import epub
        except ImportError as e:
            raise ValidationError(
                "ebooklib is required for EPUB export. Install with: pip install ebooklib"
            ) from e

        book = epub.EpubBook()

        # Metadata
        book.set_identifier(f"nexus-{content.title[:50]}-{id(content)}")
        book.set_title(content.title)
        book.set_language("en")
        book.add_author(content.author)

        # CSS
        style = epub.EpubItem(
            uid="style",
            file_name="style/default.css",
            media_type="text/css",
            content=self._epub_css(options).encode("utf-8"),
        )
        book.add_item(style)

        # Convert content to HTML chapters
        sections = content.sections or MarkdownParser.parse_sections(content.content)
        chapters = []
        spine = ["nav"]

        for i, section in enumerate(sections):
            title = section.get("title", f"Section {i + 1}")
            body_html = MarkdownParser.to_html(section.get("content", ""))

            chapter = epub.EpubHtml(
                title=title,
                file_name=f"chapter_{i + 1}.xhtml",
                lang="en",
            )
            chapter.content = f"""
            <html><head><link rel="stylesheet" href="style/default.css" /></head>
            <body>
            <h2>{title}</h2>
            {body_html}
            </body></html>
            """
            chapter.add_item(style)
            book.add_item(chapter)
            chapters.append(chapter)
            spine.append(chapter)

        # Table of contents
        book.toc = [
            epub.Link(f"chapter_{i + 1}.xhtml", ch.title, f"ch{i + 1}")
            for i, ch in enumerate(chapters)
        ]

        # Navigation
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = spine

        # Write
        buffer = io.BytesIO()
        epub.write_epub(buffer, book)
        epub_data = buffer.getvalue()

        filename = PDFExporter._safe_filename(content.title) + ".epub"

        return ExportResult(
            data=epub_data,
            filename=filename,
            mime_type="application/epub+zip",
            format="epub",
            file_size_bytes=len(epub_data),
        )

    @staticmethod
    def _epub_css(options: ExportOptions) -> str:
        """Generate EPUB stylesheet."""
        return f"""
        body {{
            font-family: Georgia, 'Times New Roman', serif;
            font-size: {options.font_size}pt;
            line-height: {options.line_spacing};
            margin: 1em;
            color: #1f2937;
        }}
        h1, h2, h3 {{
            color: {options.branding_color};
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }}
        h1 {{ font-size: 1.8em; }}
        h2 {{ font-size: 1.4em; }}
        h3 {{ font-size: 1.2em; }}
        p {{ margin-bottom: 0.8em; text-align: justify; }}
        code {{
            background-color: #f3f4f6;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: monospace;
            font-size: 0.9em;
        }}
        ul, ol {{ padding-left: 1.5em; }}
        li {{ margin-bottom: 0.3em; }}
        strong {{ color: #111827; }}
        em {{ color: #4b5563; }}
        blockquote {{
            border-left: 3px solid {options.branding_color};
            padding-left: 1em;
            margin-left: 0;
            color: #6b7280;
        }}
        """


# ── Unified Export API ───────────────────────────────────────


class ExportEngine:
    """
    Unified export engine — routes to the correct format exporter.
    """

    def __init__(self) -> None:
        self._pdf = PDFExporter()
        self._docx = DOCXExporter()
        self._epub = EPUBExporter()

    @traced("export.generate")
    async def export(
        self,
        content: ExportContent,
        options: ExportOptions | None = None,
    ) -> ExportResult:
        """Export content to the specified format."""
        if not options:
            options = ExportOptions()

        logger.info(f"Exporting '{content.title}' as {options.format}")

        if options.format == ExportFormat.PDF:
            return await self._pdf.export(content, options)
        elif options.format == ExportFormat.DOCX:
            return await self._docx.export(content, options)
        elif options.format == ExportFormat.EPUB:
            return await self._epub.export(content, options)
        elif options.format == ExportFormat.MARKDOWN:
            return self._export_markdown(content)
        elif options.format == ExportFormat.HTML:
            return self._export_html(content, options)
        elif options.format == ExportFormat.TXT:
            return self._export_text(content)
        else:
            raise ValidationError(f"Unsupported format: {options.format}")

    def _export_markdown(self, content: ExportContent) -> ExportResult:
        """Export as raw markdown."""
        md = f"# {content.title}\n\n{content.content}"
        data = md.encode("utf-8")
        return ExportResult(
            data=data,
            filename=PDFExporter._safe_filename(content.title) + ".md",
            mime_type="text/markdown",
            format="markdown",
            file_size_bytes=len(data),
        )

    def _export_html(self, content: ExportContent, options: ExportOptions) -> ExportResult:
        """Export as styled HTML."""
        body_html = MarkdownParser.to_html(content.content)
        html = f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="UTF-8">
<title>{content.title}</title>
<style>
body {{ font-family: Inter, system-ui, sans-serif; max-width: 800px;
       margin: 40px auto; padding: 20px; color: #1f2937;
       line-height: {options.line_spacing}; }}
h1 {{ color: {options.branding_color}; }} h2 {{ color: {options.branding_color}; }}
code {{ background: #f3f4f6; padding: 2px 6px; border-radius: 4px; }}
</style></head><body>
<h1>{content.title}</h1>
<p style="color:#9ca3af;font-size:0.85em">
Generated: {content.created_at or datetime.now(UTC).strftime("%B %d, %Y")}
</p><hr>{body_html}
<hr><p style="text-align:center;color:#9ca3af;font-size:0.75em">
Generated by Nexus Notebook 11 LM — ESPERANTO</p>
</body></html>"""
        data = html.encode("utf-8")
        return ExportResult(
            data=data,
            filename=PDFExporter._safe_filename(content.title) + ".html",
            mime_type="text/html",
            format="html",
            file_size_bytes=len(data),
        )

    def _export_text(self, content: ExportContent) -> ExportResult:
        """Export as plain text."""
        text = f"{content.title}\n{'=' * len(content.title)}\n\n"
        text += MarkdownParser.strip_markdown(content.content)
        data = text.encode("utf-8")
        return ExportResult(
            data=data,
            filename=PDFExporter._safe_filename(content.title) + ".txt",
            mime_type="text/plain",
            format="txt",
            file_size_bytes=len(data),
        )

    @traced("export.batch")
    async def export_batch(
        self,
        contents: list[ExportContent],
        options: ExportOptions | None = None,
    ) -> list[ExportResult]:
        """Export multiple contents to the same format."""
        results = []
        for content in contents:
            try:
                result = await self.export(content, options)
                results.append(result)
            except Exception as e:
                logger.error(f"Export failed for '{content.title}': {e}")
        return results


# Global singleton
export_engine = ExportEngine()
